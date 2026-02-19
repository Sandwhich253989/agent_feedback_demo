import sys
from pathlib import Path
from docx import Document

def main():
    title = sys.argv[1]
    content = sys.argv[2]

    output_dir = Path("/outputs")
    output_dir.mkdir(exist_ok=True)

    filepath = output_dir / "output_test.docx"

    doc = Document()
    doc.add_heading(title, level=2)

    for line in content.split("\\n"):
        doc.add_paragraph(line)

    doc.save(filepath)
    print(f"[WORKER] Saved doc to {filepath}")

if __name__ == "__main__":
    main()
