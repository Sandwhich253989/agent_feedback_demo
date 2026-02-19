from pathlib import Path
from typing import List

from docx import Document

from langchain_core.tools import tool
# from src.utils.test_file_perm import secure_file_acl



# from utils.secure_file_permissions import secure_file_acl

# --------------------TOOLS----------------------------
@tool
def write_to_doc(title: str, content: str):
    """ Writes the given content to a .docx file. """
    output_dir = Path(r"C:\Users\Public\outputs")
    output_dir.mkdir(exist_ok=True)

    filename = f"output_test2.docx"
    filepath = output_dir / filename

    doc = Document()

    doc.add_heading(f"{title}\n", level=2)
    for line in content.split("\n"):
        line = line.strip()
        doc.add_paragraph(line)

    # print(filepath)
    print(f"[TOOL]: Tool write_to_doc executed: Word document saved to  f{filepath}")
    doc.save(filepath)

    # secure_file_acl(filepath.name,"test")
    return f"Word document saved to {filepath}"


@tool
def write_sections_to_doc(title: str, sections: List[dict]):
    """
    Writes multiple sections to a .docx file with metadata.
    Each section dict should have: {name, content, confidence, status}
    """
    output_dir = Path(r"C:\Users\Public\outputs")
    output_dir.mkdir(exist_ok=True)

    filename = f"output_test2.docx"
    filepath = output_dir / filename

    doc = Document()

    # Main title
    doc.add_heading(title, level=1)

    # Add metadata paragraph
    total_sections = len(sections)
    auto_approved = sum(1 for s in sections if s.get("status") == "auto_approved")
    human_reviewed = sum(1 for s in sections if s.get("status") == "human_reviewed")

    metadata = doc.add_paragraph()
    metadata.add_run(f"Document Statistics:\n").bold = True
    metadata.add_run(f"Total Sections: {total_sections}\n")
    metadata.add_run(f"Auto-approved: {auto_approved}\n")
    metadata.add_run(f"Human-reviewed: {human_reviewed}\n")
    metadata.add_run(f"Efficiency: {(auto_approved/total_sections*100):.1f}% automated\n\n")

    # Add each section
    for section in sections:
        section_name = section.get("name", "Untitled Section")
        content = section.get("content", "")
        confidence = section.get("confidence", 0.0)
        status = section.get("status", "unknown")

        # Section heading with status indicator
        heading = doc.add_heading(section_name, level=2)

        # Add status badge
        status_para = doc.add_paragraph()
        if status == "auto_approved":
            status_para.add_run(f"Auto-approved (Confidence: {confidence:.2f})").italic = True
        elif status == "human_reviewed":
            status_para.add_run(f"Human-reviewed (Confidence: {confidence:.2f})").italic = True
        else:
            status_para.add_run(f"Status: {status} (Confidence: {confidence:.2f})").italic = True

        # Add content
        for line in content.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)

        # Add spacing
        doc.add_paragraph()

    # assign File permissions to TARGET USER
    # secure_file_acl(str(filepath),"hp")
    doc.save(filepath)
    print(f"[TOOL]: Sections written to {filepath}")
    return f"Document with {total_sections} sections saved to {filepath}"

